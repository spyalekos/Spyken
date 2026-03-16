import flet as ft
import fitz  # PyMuPDF
import docx
import edge_tts
import asyncio
import os
import tempfile
import textwrap

VOICE_MALE = "el-GR-NestorasNeural"
VOICE_FEMALE = "el-GR-AthinaNeural"
VOICE_EN_MALE = "en-GB-RyanNeural"
VOICE_EN_FEMALE = "en-GB-SoniaNeural"

# Video frame resolution
VIDEO_W = 1280
VIDEO_H = 720
VIDEO_FPS = 10  # Reduced from 24 to speed up rendering significantly

# Fluorescent green highlight color
HIGHLIGHT_COLOR = (57, 255, 20)


def is_english(text: str) -> bool:
    """Returns True if the majority of letters in text are Latin (English)."""
    latin = sum(1 for c in text if 'A' <= c <= 'Z' or 'a' <= c <= 'z')
    greek = sum(1 for c in text if '\u0370' <= c <= '\u03ff' or '\u1f00' <= c <= '\u1fff')
    return latin > greek


def is_valid_text(text: str) -> bool:
    """Returns True only if text has enough real words for TTS."""
    if len(text) < 3:
        return False
    if not any(c.isalpha() for c in text):
        return False
    return True


import re as _re
_EMOJI_RE = _re.compile(
    "["
    "\U0001F000-\U0001FFFF"  # emoji, symbols, flags
    "\U00002300-\U000027BF"  # misc technical/dingbats
    "\U0000FE00-\U0000FE0F"  # variation selectors
    "\U0001FA00-\U0001FAFF"  # extended symbols
    "]+", flags=_re.UNICODE
)

def clean_for_tts(text: str) -> str:
    """
    Remove emoji and symbol characters that cause edge_tts to produce
    garbled word-boundary events (character spans instead of words).
    The cleaned text is used ONLY for TTS; original text is kept for
    PDF word-rect matching.
    """
    cleaned = _EMOJI_RE.sub(' ', text)
    cleaned = _re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


def chunk_text(text: str, max_chars: int = 800) -> list[str]:
    """Split text into chunks of max_chars, breaking at sentence boundaries."""
    if len(text) <= max_chars:
        return [text]
    chunks = []
    while len(text) > max_chars:
        split_pos = max_chars
        for sep in ('. ', '! ', '? ', ', '):
            pos = text.rfind(sep, 0, max_chars)
            if pos != -1:
                split_pos = pos + len(sep)
                break
        chunks.append(text[:split_pos].strip())
        text = text[split_pos:].strip()
    if text:
        chunks.append(text)
    return chunks


def extract_paragraphs(filepath: str) -> list[str]:
    ext = filepath.lower().split('.')[-1]
    paragraphs = []

    if ext == 'docx':
        doc = docx.Document(filepath)
        for p in doc.paragraphs:
            text = p.text.strip()
            if is_valid_text(text):
                paragraphs.append(text)
    elif ext == 'pdf':
        doc = fitz.open(filepath)
        for page in doc:
            blocks = page.get_text("blocks")
            merged = merge_pdf_blocks(blocks)
            for (x0, y0, x1, y1, text) in merged:
                if is_valid_text(text):
                    paragraphs.append(text)
    else:
        raise ValueError("Μη υποστηριζόμενη μορφή αρχείου")

    return paragraphs


def merge_pdf_blocks(blocks: list) -> list:
    """
    Merge adjacent PDF text blocks that belong to the same paragraph.
    Two consecutive blocks are merged when they are vertically close
    (gap ≤ 1.8 × estimated line height) and horizontally overlapping.

    Input : list of fitz block tuples (x0, y0, x1, y1, text, block_no, block_type)
    Returns: list of (x0, y0, x1, y1, text) tuples.
    """
    # Keep only text blocks (block_type == 0)
    text_blocks = [b for b in blocks if b[6] == 0]

    if not text_blocks:
        return []

    merged = []
    cur_x0, cur_y0, cur_x1, cur_y1, cur_text = (
        text_blocks[0][0], text_blocks[0][1],
        text_blocks[0][2], text_blocks[0][3],
        text_blocks[0][4].strip(),
    )

    for b in text_blocks[1:]:
        bx0, by0, bx1, by1, btxt = b[0], b[1], b[2], b[3], b[4].strip()

        # Estimate line height of current block
        cur_height = cur_y1 - cur_y0
        lines_estimate = max(1, cur_text.count('\n') + 1)
        line_h = cur_height / lines_estimate

        # Vertical gap between end of current block and start of next
        v_gap = by0 - cur_y1

        # Horizontal overlap check: blocks share at least some horizontal range
        h_overlap = min(cur_x1, bx1) - max(cur_x0, bx0)

        if v_gap >= 0 and v_gap <= line_h * 1.8 and h_overlap > -20:
            # Merge: join with a space (replace internal newlines too)
            cur_text = cur_text + " " + btxt
            cur_x0 = min(cur_x0, bx0)
            cur_y1 = by1
            cur_x1 = max(cur_x1, bx1)
        else:
            merged.append((cur_x0, cur_y0, cur_x1, cur_y1, cur_text))
            cur_x0, cur_y0, cur_x1, cur_y1, cur_text = bx0, by0, bx1, by1, btxt

    merged.append((cur_x0, cur_y0, cur_x1, cur_y1, cur_text))

    # Clean up internal newlines in each merged block
    result = []
    for (x0, y0, x1, y1, t) in merged:
        clean = " ".join(t.splitlines()).strip()
        result.append((x0, y0, x1, y1, clean))

    return result


# ─────────────────────────── TTS WITH WORD TIMING ─────────────────────────────

async def generate_tts_with_word_timings(text: str, voice: str, out_path: str) -> list[dict]:
    """
    Stream TTS audio and collect WordBoundary events.
    Uses clean_for_tts(text) to strip emoji before sending to edge_tts,
    so word-boundary events contain proper words instead of character spans.
    """
    tts_text = clean_for_tts(text)
    if not tts_text:
        return []
    word_timings = []
    audio_bytes = bytearray()

    for attempt in range(3):
        try:
            communicate = edge_tts.Communicate(tts_text, voice, boundary="WordBoundary")
            word_timings.clear()
            audio_bytes.clear()

            async for chunk in communicate.stream():
                if chunk["type"] == "audio":
                    audio_bytes.extend(chunk["data"])
                elif chunk["type"] == "WordBoundary":
                    # offset and duration are in 100-nanosecond units
                    offset_s = chunk["offset"] / 1e7
                    duration_s = chunk["duration"] / 1e7
                    word = chunk.get("text", "")
                    word_timings.append({
                        "offset_s": offset_s,
                        "duration_s": duration_s,
                        "word": word,
                    })

            if audio_bytes:
                with open(out_path, "wb") as f:
                    f.write(audio_bytes)
                return word_timings

        except Exception:
            await asyncio.sleep(0.5)

    return []


async def generate_tts_chunk(text: str, voice: str, out_path: str) -> bool:
    """Generate a single TTS mp3 chunk (audio only). Returns True on success."""
    tts_text = clean_for_tts(text)
    if not tts_text:
        return False
    for attempt in range(3):
        try:
            communicate = edge_tts.Communicate(tts_text, voice)
            await communicate.save(out_path)
            if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
                return True
        except Exception:
            await asyncio.sleep(0.5)
    return False


# ─────────────────────────── VIDEO HELPERS ────────────────────────────────────

def extract_paragraphs_pdf_with_pos(filepath: str) -> list[tuple]:
    """
    Returns list of (text, page_idx, fitz.Rect) for each valid paragraph.
    """
    results = []
    doc = fitz.open(filepath)
    for page_idx, page in enumerate(doc):
        blocks = page.get_text("blocks")
        merged = merge_pdf_blocks(blocks)
        for (x0, y0, x1, y1, text) in merged:
            if is_valid_text(text):
                rect = fitz.Rect(x0, y0, x1, y1)
                results.append((text, page_idx, rect, doc))
    return results


def get_pdf_word_rects(pdf_doc, page_idx: int, para_rect: fitz.Rect) -> list[tuple]:
    """
    Return list of (word_text, fitz.Rect) for words within para_rect on page_idx.
    Uses fitz word-level extraction.
    """
    page = pdf_doc[page_idx]
    # get_text("words") returns (x0, y0, x1, y1, "word", block_no, line_no, word_no)
    words = page.get_text("words")
    result = []
    for w in words:
        word_rect = fitz.Rect(w[0], w[1], w[2], w[3])
        word_text = w[4]
        # Skip words that are pure emoji / symbols (TTS won't produce word-boundaries for them)
        if not any(c.isalpha() or c.isdigit() for c in word_text):
            continue
        # Only include words that intersect with the paragraph block
        if word_rect.intersects(para_rect):
            result.append((word_text, word_rect))
    return result


def render_page_pdf_image(
    pdf_doc,
    page_idx: int,
    highlight_rect=None,
    word_highlight_rect=None,
    target_w: int = VIDEO_W,
    target_h: int = VIDEO_H,
):
    """
    Render a PDF page as a PIL RGBA image sized to fit target_w x target_h.
    - highlight_rect: paragraph-level semi-transparent highlight (fitz.Rect)
    - word_highlight_rect: word-level bright highlight (fitz.Rect), drawn on top
    Returns a PIL Image (RGB).
    """
    from PIL import Image, ImageDraw

    page = pdf_doc[page_idx]
    page_rect = page.rect
    scale = min(target_w / page_rect.width, target_h / page_rect.height)
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False)

    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    # Centre the page on a dark canvas
    canvas = Image.new("RGB", (target_w, target_h), (30, 30, 40))
    x_off = (target_w - img.width) // 2
    y_off = (target_h - img.height) // 2
    canvas.paste(img, (x_off, y_off))

    overlay = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)

    def scaled_rect(r):
        return (
            int(r.x0 * scale) + x_off,
            int(r.y0 * scale) + y_off,
            int(r.x1 * scale) + x_off,
            int(r.y1 * scale) + y_off,
        )

    # Paragraph-level backdrop highlight (very faint)
    if highlight_rect is not None:
        hx0, hy0, hx1, hy1 = scaled_rect(highlight_rect)
        draw.rectangle([hx0, hy0, hx1, hy1], fill=(57, 255, 20, 30))
        draw.rectangle([hx0, hy0, hx1, hy1], outline=(57, 255, 20, 80), width=2)

    # Word-level highlight (bright on top)
    if word_highlight_rect is not None:
        wx0, wy0, wx1, wy1 = scaled_rect(word_highlight_rect)
        draw.rectangle([wx0, wy0, wx1, wy1], fill=(57, 255, 20, 120))
        draw.rectangle([wx0, wy0, wx1, wy1], outline=(57, 255, 20, 255), width=3)

    canvas = canvas.convert("RGBA")
    canvas = Image.alpha_composite(canvas, overlay).convert("RGB")

    return canvas


def _build_docx_layout(text: str, box_w: int, box_h: int):
    """
    Compute line wrapping and best font size for DOCX renderer.
    Returns (best_lines, best_font, line_h, font_size).
    """
    from PIL import ImageFont

    def try_wrap(font_sz):
        try:
            font = ImageFont.truetype("arial.ttf", size=font_sz)
        except Exception:
            font = ImageFont.load_default()
            font_sz = 14

        words = text.split()
        if not words:
            return ["(κενή παράγραφος)"], font, font_sz

        lines = []
        current_line = []
        for word in words:
            test_line = " ".join(current_line + [word])
            w = font.getbbox(test_line)[2] if hasattr(font, 'getbbox') else font.getsize(test_line)[0]
            if w <= box_w:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(" ".join(current_line))
                    current_line = [word]
                else:
                    lines.append(word)
                    current_line = []
        if current_line:
            lines.append(" ".join(current_line))

        return lines, font, font_sz

    best_lines = []
    best_font = None
    best_font_size = 50
    line_h = 50

    for f_size in range(50, 15, -2):
        lines, font, actual_f_size = try_wrap(f_size)
        test_line_h = max(22, int(actual_f_size * 1.5))
        total_h = len(lines) * test_line_h

        if total_h <= box_h or f_size == 16:
            best_lines = lines
            best_font = font
            best_font_size = actual_f_size
            line_h = test_line_h
            break

    if not best_lines:
        best_lines, best_font, best_font_size = try_wrap(24)
        line_h = 36

    return best_lines, best_font, line_h, best_font_size


def render_docx_paragraph_image(
    text: str,
    para_idx: int,
    total: int,
    highlight_word: str = None,
    target_w: int = VIDEO_W,
    target_h: int = VIDEO_H,
):
    """
    Render a DOCX paragraph as a PIL Image.
    - If highlight_word is given, that specific word (first occurrence on screen) is
      highlighted in bright fluorescent green; the rest of the text box has a dimmer backdrop.
    - Otherwise falls back to the full-paragraph highlight.
    Returns a PIL Image (RGB).
    """
    from PIL import Image, ImageDraw, ImageFont

    BG = (24, 28, 42)
    HEADER_BG = (40, 44, 64)
    H_COLOR = HIGHLIGHT_COLOR          # (57, 255, 20)
    TEXT_COLOR = (30, 30, 30)
    HEADER_COLOR = (180, 180, 200)

    canvas = Image.new("RGB", (target_w, target_h), BG)
    draw = ImageDraw.Draw(canvas)

    # Header bar
    draw.rectangle([0, 0, target_w, 60], fill=HEADER_BG)
    header_txt = f"Παράγραφος {para_idx + 1} / {total}"
    try:
        header_font = ImageFont.truetype("arial.ttf", size=24)
        draw.text((target_w // 2, 30), header_txt, font=header_font, fill=HEADER_COLOR, anchor="mm")
    except Exception:
        draw.text((target_w // 2, 30), header_txt, fill=HEADER_COLOR, anchor="mm")

    # Layout dimensions
    box_w = target_w - 120
    box_h = target_h - 160
    box_x = 60
    box_y = 90

    best_lines, best_font, line_h, font_size = _build_docx_layout(text, box_w, box_h)

    # Compute overall text block dimensions
    text_block_h = len(best_lines) * line_h + 40
    max_line_w = 0
    for line in best_lines:
        w = best_font.getbbox(line)[2] if hasattr(best_font, 'getbbox') else best_font.getsize(line)[0]
        max_line_w = max(max_line_w, w)
    text_block_w = min(max_line_w + 40, box_w + 40)

    tx = box_x - 20
    ty = box_y - 10

    if highlight_word:
        # Draw a dim backdrop for the whole paragraph box
        draw.rectangle([tx, ty, tx + text_block_w, ty + text_block_h],
                       fill=(30, 80, 30))
        draw.rectangle([tx, ty, tx + text_block_w, ty + text_block_h],
                       outline=(0, 120, 0), width=2)
    else:
        # Full bright highlight (original behaviour)
        draw.rectangle([tx, ty, tx + text_block_w, ty + text_block_h], fill=H_COLOR)
        draw.rectangle([tx, ty, tx + text_block_w, ty + text_block_h], outline=(0, 200, 0), width=4)

    # Draw text lines and find word bounding box
    y_cursor = box_y + 10
    word_rect_found = None  # (x0, y0, x1, y1) on canvas

    # Normalize highlight_word for comparison (strip punctuation)
    def normalize(w):
        return w.strip(".,;:!?\"'()[]»«—–-").lower() if w else ""

    norm_hw = normalize(highlight_word)
    word_found = False

    for line in best_lines:
        line_words = line.split()
        x_cursor = box_x + 10

        for lw in line_words:
            # Measure this word
            if hasattr(best_font, 'getbbox'):
                wb = best_font.getbbox(lw)
                w_width = wb[2] - wb[0]
                w_height = wb[3] - wb[1]
            else:
                w_width, w_height = best_font.getsize(lw)
                wb = (0, 0, w_width, w_height)

            # Draw the word
            if highlight_word and not word_found and normalize(lw) == norm_hw:
                # Highlight this word
                pad = 4
                hx0 = x_cursor - pad
                hy0 = y_cursor - pad
                hx1 = x_cursor + w_width + pad
                hy1 = y_cursor + line_h - 2
                draw.rectangle([hx0, hy0, hx1, hy1], fill=H_COLOR)
                draw.rectangle([hx0, hy0, hx1, hy1], outline=(0, 220, 0), width=2)
                draw.text((x_cursor, y_cursor), lw, font=best_font, fill=TEXT_COLOR)
                word_found = True
            else:
                # Normal text (white-ish on dark backdrop when word mode)
                txt_col = (200, 220, 200) if highlight_word else TEXT_COLOR
                draw.text((x_cursor, y_cursor), lw, font=best_font, fill=txt_col)

            # Advance x by word width + space width
            if hasattr(best_font, 'getbbox'):
                space_w = best_font.getbbox(" ")[2]
            else:
                space_w = best_font.getsize(" ")[0]
            x_cursor += w_width + space_w

        y_cursor += line_h
        if y_cursor > box_y + box_h:
            break

    # Footer
    draw.rectangle([0, target_h - 40, target_w, target_h], fill=HEADER_BG)
    footer_txt = "Spyken · MP4 by spyalekos"
    try:
        footer_font = ImageFont.truetype("arial.ttf", size=18)
        draw.text((target_w // 2, target_h - 20), footer_txt,
                  font=footer_font, fill=(100, 100, 130), anchor="mm")
    except Exception:
        draw.text((target_w // 2, target_h - 20), footer_txt, fill=(100, 100, 130), anchor="mm")

    return canvas


def get_mp3_duration(path: str) -> float:
    """Return duration of an mp3 file in seconds."""
    try:
        from mutagen.mp3 import MP3
        audio = MP3(path)
        return audio.info.length
    except Exception:
        return 3.0  # fallback


# ─────────────────────── WORD-TIMING ALIGNMENT ────────────────────────────────

def align_word_timings_to_text(word_timings: list[dict], text: str) -> list[dict]:
    """
    The TTS engine may return WordBoundary words in a slightly different order
    or with punctuation stripped. We do a best-effort match of timing words
    to the actual words in `text` so that highlights correspond to visible tokens.

    Returns same list but with an extra 'text_word' key containing the original
    word from the text (or same as 'word' if no match found).
    """
    text_words = text.split()
    timing_words = [t.copy() for t in word_timings]

    t_idx = 0
    for tw in timing_words:
        # Advance text_words pointer to find the best match
        norm_tw = tw["word"].strip(".,;:!?\"'()[]»«—–-").lower()
        matched = False
        for offset in range(min(12, len(text_words) - t_idx)):
            candidate = text_words[t_idx + offset].strip(".,;:!?\"'()[]»«—–-").lower()
            if candidate == norm_tw or (
                norm_tw and candidate and (norm_tw in candidate or candidate in norm_tw)
            ):
                tw["text_word"] = text_words[t_idx + offset]
                t_idx = t_idx + offset + 1
                matched = True
                break
        if not matched:
            tw["text_word"] = tw["word"]

    return timing_words


async def convert_to_video(filepath: str, output_path: str, progress_callback):
    """
    Master function: extract paragraphs (with positions for PDF),
    generate TTS per paragraph with word timings, build per-word video frames,
    assemble mp4.
    """
    import numpy as np
    from moviepy import ImageClip, AudioFileClip, concatenate_videoclips, concatenate_audioclips
    from proglog import ProgressBarLogger

    class FletMoviepyLogger(ProgressBarLogger):
        def __init__(self, cb):
            super().__init__()
            self.ui_callback = cb

        def bars_callback(self, bar, attr, value, old_value=None):
            if attr == 'index' and bar in ('chunk', 't', 'frame_index'):
                total = self.bars[bar].get('total', 1)
                if total > 0:
                    pct = int((value / total) * 100)
                    msg = "Μίξη ήχου..." if bar == 'chunk' else "Συναρμολόγηση βίντεο..."
                    self.ui_callback(value, total, f"{msg} {pct}%")

    ext = filepath.lower().split('.')[-1]
    temp_dir = tempfile.mkdtemp()

    try:
        # ── 1. Extract paragraphs ─────────────────────────────────────────────
        if ext == 'pdf':
            pdf_doc = fitz.open(filepath)
            para_data = []  # (text, page_idx, rect)
            for page_idx, page in enumerate(pdf_doc):
                blocks = page.get_text("blocks")
                merged = merge_pdf_blocks(blocks)
                for (x0, y0, x1, y1, text) in merged:
                    if is_valid_text(text):
                        rect = fitz.Rect(x0, y0, x1, y1)
                        para_data.append((text, page_idx, rect))
        else:  # docx
            doc_obj = docx.Document(filepath)
            para_data = []
            for p in doc_obj.paragraphs:
                text = p.text.strip()
                if is_valid_text(text):
                    para_data.append((text, len(para_data), None))

        if not para_data:
            raise ValueError("Δεν βρέθηκαν παράγραφοι στο αρχείο.")

        total = len(para_data)
        clips = []
        voice_index = 0

        # ── 2. Per-paragraph: TTS with word timings + frames ──────────────────
        for i, para_item in enumerate(para_data):
            text = para_item[0]
            progress_callback(i, total, f"Παράγραφος {i+1}/{total}: TTS + λέξεις…")
            await asyncio.sleep(0)  # yield to UI

            # Pick voice
            voice = (
                (VOICE_EN_MALE if voice_index % 2 == 0 else VOICE_EN_FEMALE)
                if is_english(text)
                else (VOICE_MALE if voice_index % 2 == 0 else VOICE_FEMALE)
            )

            # For long paragraphs we chunk the text
            chunks = chunk_text(text, 800)
            chunk_audio_clips = []
            all_word_timings = []     # accumulated across chunks
            chunk_time_offset = 0.0  # running time offset for multi-chunk paragraphs

            for c_idx, chunk in enumerate(chunks):
                chunk_audio_path = os.path.join(temp_dir, f"audio_{i}_{c_idx}.mp3")
                word_timings = await generate_tts_with_word_timings(chunk, voice, chunk_audio_path)

                if os.path.exists(chunk_audio_path) and os.path.getsize(chunk_audio_path) > 0:
                    ac = AudioFileClip(chunk_audio_path)
                    chunk_audio_clips.append(ac)

                    # Shift word timings by the running offset
                    for wt in word_timings:
                        shifted = dict(wt)
                        shifted["offset_s"] += chunk_time_offset
                        all_word_timings.append(shifted)

                    chunk_time_offset += ac.duration
                else:
                    # Fallback: plain TTS without timings
                    ok = await generate_tts_chunk(chunk, voice, chunk_audio_path)
                    if ok:
                        ac = AudioFileClip(chunk_audio_path)
                        chunk_audio_clips.append(ac)
                        chunk_time_offset += ac.duration

            if chunk_audio_clips:
                if len(chunk_audio_clips) == 1:
                    combined_audio = chunk_audio_clips[0]
                else:
                    combined_audio = concatenate_audioclips(chunk_audio_clips)
                total_duration = combined_audio.duration
                voice_index += 1
            else:
                combined_audio = None
                total_duration = 3.0

            # ── 3. Align word timings to actual text ──────────────────────────
            if all_word_timings:
                all_word_timings = align_word_timings_to_text(all_word_timings, text)

            # ── 4. Build per-word frames ───────────────────────────────────────
            progress_callback(i, total, f"Παράγραφος {i+1}/{total}: Frames…")
            await asyncio.sleep(0)

            # Pre-fetch PDF word rects (only for PDF with word timings)
            pdf_word_rects = []
            if ext == 'pdf' and all_word_timings:
                _, page_idx, para_rect = para_item
                pdf_word_rects = get_pdf_word_rects(pdf_doc, page_idx, para_rect)

            para_clips = []

            if all_word_timings:
                # ── Word-level frame generation ───────────────────────────────

                # For PDF: pre-render the base page image ONCE (no highlight),
                # then composite word highlights on top.
                # This avoids re-rendering the full page for every word.
                if ext == 'pdf':
                    _, page_idx, para_rect = para_item
                    base_pdf_img = render_page_pdf_image(
                        pdf_doc,
                        page_idx,
                        highlight_rect=para_rect,
                        word_highlight_rect=None,
                    )
                    # Normalise word rects for this page (scale + offset)
                    page = pdf_doc[page_idx]
                    page_rect = page.rect
                    scale = min(VIDEO_W / page_rect.width, VIDEO_H / page_rect.height)
                    x_off = (VIDEO_W - int(page_rect.width * scale)) // 2
                    y_off = (VIDEO_H - int(page_rect.height * scale)) // 2

                    # Sequential pointer into pdf_word_rects
                    wr_ptr = 0

                    # Pre-roll: blank (no word highlight) frame before first word
                    first_offset = all_word_timings[0]["offset_s"]
                    if first_offset > 0.05:
                        para_clips.append(ImageClip(np.array(base_pdf_img), duration=first_offset))

                else:
                    # DOCX pre-roll
                    _, para_idx, _ = para_item
                    first_offset = all_word_timings[0]["offset_s"]
                    if first_offset > 0.05:
                        pre_frame = render_docx_paragraph_image(text, para_idx, total)
                        para_clips.append(ImageClip(np.array(pre_frame), duration=first_offset))

                n_timings = len(all_word_timings)
                for w_idx, wt in enumerate(all_word_timings):
                    # Clip duration = gap to next word's offset (covers silence between words)
                    if w_idx < n_timings - 1:
                        clip_dur = all_word_timings[w_idx + 1]["offset_s"] - wt["offset_s"]
                    else:
                        clip_dur = total_duration - wt["offset_s"]
                    if clip_dur < 0.04:
                        clip_dur = 0.04

                    word_text = wt.get("text_word", wt["word"])

                    if ext == 'pdf':
                        # Sequential search: advance wr_ptr instead of restarting
                        norm_word = word_text.strip(".,;:!?\"'()[]\u00bb\u00ab\u2014\u2013-").lower()
                        matching_rect = None
                        search_limit = min(wr_ptr + 20, len(pdf_word_rects))
                        for j in range(wr_ptr, search_limit):
                            wr_text, wr_rect = pdf_word_rects[j]
                            wr_norm = wr_text.strip(".,;:!?\"'()[]\u00bb\u00ab\u2014\u2013-").lower()
                            if wr_norm == norm_word or (
                                wr_norm and norm_word and
                                (norm_word in wr_norm or wr_norm in norm_word)
                            ):
                                matching_rect = wr_rect
                                wr_ptr = j + 1  # advance past this word
                                break
                        else:
                            # No match found — advance ptr by 1 to avoid stalling
                            if wr_ptr < len(pdf_word_rects):
                                wr_ptr += 1

                        # Composite highlight onto cached base image
                        from PIL import Image, ImageDraw
                        frame_img = base_pdf_img.copy()
                        if matching_rect is not None:
                            overlay = Image.new("RGBA", (VIDEO_W, VIDEO_H), (0, 0, 0, 0))
                            draw = ImageDraw.Draw(overlay)
                            wx0 = int(matching_rect.x0 * scale) + x_off
                            wy0 = int(matching_rect.y0 * scale) + y_off
                            wx1 = int(matching_rect.x1 * scale) + x_off
                            wy1 = int(matching_rect.y1 * scale) + y_off
                            draw.rectangle([wx0, wy0, wx1, wy1], fill=(57, 255, 20, 140))
                            draw.rectangle([wx0, wy0, wx1, wy1],
                                           outline=(57, 255, 20, 255), width=3)
                            frame_img = Image.alpha_composite(
                                frame_img.convert("RGBA"), overlay
                            ).convert("RGB")
                    else:
                        _, para_idx, _ = para_item
                        frame_img = render_docx_paragraph_image(
                            text, para_idx, total,
                            highlight_word=word_text,
                        )

                    frame_np = np.array(frame_img)
                    para_clips.append(ImageClip(frame_np, duration=clip_dur))

                # No gap-fill needed: offset-based durations already cover total_duration

            else:
                # ── Fallback: paragraph-level (original behaviour) ────────────
                if ext == 'pdf':
                    _, page_idx, rect = para_item
                    frame_img = render_page_pdf_image(pdf_doc, page_idx, highlight_rect=rect)
                else:
                    _, para_idx, _ = para_item
                    frame_img = render_docx_paragraph_image(text, para_idx, total)

                frame_np = np.array(frame_img)
                para_clips.append(ImageClip(frame_np, duration=total_duration if total_duration else 3.0))

            # ── 5. Concatenate word clips → paragraph clip ─────────────────────
            if len(para_clips) == 1:
                para_video = para_clips[0]
            else:
                para_video = concatenate_videoclips(para_clips, method="compose")

            if combined_audio:
                audio_dur = combined_audio.duration
                video_dur = para_video.duration
                if audio_dur > video_dur:
                    combined_audio = combined_audio.subclipped(0, video_dur)
                para_video = para_video.with_audio(combined_audio)

            clips.append(para_video)

        # ── 6. Assemble final video ────────────────────────────────────────────
        progress_callback(total, total, "Συναρμολόγηση βίντεο...")
        await asyncio.sleep(0)

        ui_logger = FletMoviepyLogger(progress_callback)

        final = concatenate_videoclips(clips, method="compose")

        def _write_video():
            final.write_videofile(
                output_path,
                fps=VIDEO_FPS,
                codec="libx264",
                audio_codec="aac",
                temp_audiofile=os.path.join(temp_dir, "tmp_audio.m4a"),
                remove_temp=True,
                logger=ui_logger,
            )
            final.close()

        await asyncio.to_thread(_write_video)

    finally:
        # Cleanup temp files
        for f in os.listdir(temp_dir):
            try:
                os.remove(os.path.join(temp_dir, f))
            except Exception:
                pass
        try:
            os.rmdir(temp_dir)
        except Exception:
            pass


# ──────────────────────────── AUDIO CONVERSION ────────────────────────────────

async def convert_to_audio(paragraphs: list[str], output_path: str, progress_callback):
    temp_dir = tempfile.mkdtemp()
    temp_files = []

    all_chunks = []
    for p in paragraphs:
        all_chunks.extend(chunk_text(p))

    total = len(all_chunks)
    voice_index = 0

    for i, chunk in enumerate(all_chunks):
        voice = (
            (VOICE_EN_MALE if voice_index % 2 == 0 else VOICE_EN_FEMALE)
            if is_english(chunk)
            else (VOICE_MALE if voice_index % 2 == 0 else VOICE_FEMALE)
        )
        temp_file = os.path.join(temp_dir, f"part_{i}.mp3")

        success = False
        for attempt in range(3):
            try:
                communicate = edge_tts.Communicate(chunk, voice)
                await communicate.save(temp_file)
                if os.path.exists(temp_file) and os.path.getsize(temp_file) > 0:
                    success = True
                    break
            except Exception:
                await asyncio.sleep(0.5)

        if success:
            temp_files.append(temp_file)
            voice_index += 1

        progress_callback(i + 1, total)

    with open(output_path, "wb") as outfile:
        for tf in temp_files:
            with open(tf, "rb") as infile:
                outfile.write(infile.read())

    for tf in temp_files:
        try:
            os.remove(tf)
        except Exception:
            pass
    try:
        os.rmdir(temp_dir)
    except Exception:
        pass


# ──────────────────────────────── UI ──────────────────────────────────────────

def main(page: ft.Page):
    APP_VERSION = "1.5.0"
    page.title = "Spyken by spyalekos - Έγγραφο σε Ομιλία (MP3) & Βίντεο (MP4)"
    page.window.width = 680
    page.window.height = 740
    page.theme_mode = ft.ThemeMode.DARK
    page.padding = 30
    page.window.center()

    # UI Elements
    selected_files_list = ft.ListView(expand=1, spacing=10, padding=20, auto_scroll=True)
    status_text = ft.Text("Κατάσταση: Σε αναμονή", size=16, color=ft.Colors.GREY_400)
    progress_bar = ft.ProgressBar(width=440, color="amber", bgcolor="#263238", value=0)
    progress_bar.visible = False

    file_queue = []

    def log(msg, error=False):
        color = ft.Colors.RED_400 if error else ft.Colors.GREEN_400
        selected_files_list.controls.append(ft.Text(msg, color=color))
        page.update()

    def set_all_buttons(disabled: bool):
        pick_btn.disabled = disabled
        clear_btn.disabled = disabled
        convert_btn.disabled = disabled
        video_btn.disabled = disabled
        page.update()

    async def pick_files_clicked(e):
        file_picker = ft.FilePicker()
        files_list = await file_picker.pick_files(allow_multiple=True, allowed_extensions=["pdf", "docx"])
        if files_list:
            for f in files_list:
                if f.path not in file_queue:
                    file_queue.append(f.path)
                    selected_files_list.controls.append(ft.Text(f"Επιλέχθηκε: {f.name}", color=ft.Colors.WHITE70))
            page.update()

    def clear_files(e):
        file_queue.clear()
        selected_files_list.controls.clear()
        selected_files_list.controls.append(ft.Text("Η λίστα καθαρίστηκε", color=ft.Colors.GREY_400))
        page.update()

    # ── MP3 Conversion ────────────────────────────────────────────────────────

    async def start_conversion(e):
        if not file_queue:
            status_text.value = "Παρακαλώ επιλέξτε αρχεία πρώτα!"
            status_text.color = ft.Colors.RED_400
            page.update()
            return

        set_all_buttons(True)
        progress_bar.visible = True
        progress_bar.value = 0
        page.update()

        for filepath in file_queue:
            status_text.value = f"Εξαγωγή κειμένου: {os.path.basename(filepath)}"
            status_text.color = ft.Colors.BLUE_400
            page.update()

            try:
                await asyncio.sleep(0.1)
                paragraphs = extract_paragraphs(filepath)
                if not paragraphs:
                    log(f"Δεν βρέθηκε κείμενο στο {os.path.basename(filepath)}", error=True)
                    continue

                output_path = os.path.splitext(filepath)[0] + ".mp3"

                def update_progress(current, total):
                    progress_bar.value = current / total
                    status_text.value = f"Δημιουργία ήχου: {current}/{total} παράγραφοι"
                    page.update()

                await convert_to_audio(paragraphs, output_path, update_progress)
                log(f"Ολοκληρώθηκε: {os.path.basename(output_path)}")

            except Exception as ex:
                log(f"Σφάλμα στο {os.path.basename(filepath)}: {str(ex)}", error=True)

        status_text.value = "Κατάσταση: Όλες οι μετατροπές ολοκληρώθηκαν!"
        status_text.color = ft.Colors.GREEN_400
        progress_bar.visible = False
        set_all_buttons(False)

    # ── MP4 Conversion ────────────────────────────────────────────────────────

    async def start_video_conversion(e):
        if not file_queue:
            status_text.value = "Παρακαλώ επιλέξτε αρχεία πρώτα!"
            status_text.color = ft.Colors.RED_400
            page.update()
            return

        set_all_buttons(True)
        progress_bar.visible = True
        progress_bar.value = 0
        page.update()

        for filepath in file_queue:
            ext = filepath.lower().split('.')[-1]
            if ext not in ('pdf', 'docx'):
                log(f"Μη υποστηριζόμενο αρχείο για βίντεο: {os.path.basename(filepath)}", error=True)
                continue

            status_text.value = f"🎬 Ξεκινά δημιουργία βίντεο: {os.path.basename(filepath)}"
            status_text.color = ft.Colors.PURPLE_300
            page.update()

            try:
                output_path = os.path.splitext(filepath)[0] + ".mp4"
                loop = asyncio.get_running_loop()

                def update_video_progress(current, total, msg=""):
                    def _update_ui():
                        progress_bar.value = current / max(total, 1)
                        status_text.value = f"🎬 {msg}" if msg else f"🎬 Παράγραφος {current}/{total}"
                        status_text.color = ft.Colors.PURPLE_300
                        page.update()
                    loop.call_soon_threadsafe(_update_ui)

                await convert_to_video(filepath, output_path, update_video_progress)
                log(f"🎬 Βίντεο: {os.path.basename(output_path)}")

            except Exception as ex:
                log(f"Σφάλμα βίντεο στο {os.path.basename(filepath)}: {str(ex)}", error=True)

        status_text.value = "Κατάσταση: Δημιουργία βίντεο ολοκληρώθηκε!"
        status_text.color = ft.Colors.GREEN_400
        progress_bar.visible = False
        set_all_buttons(False)

    # ── About dialog ──────────────────────────────────────────────────────────

    def open_about(e):
        dlg = ft.AlertDialog(
            title=ft.Text("🎵 Spyken", size=22, weight=ft.FontWeight.BOLD, color=ft.Colors.AMBER_400),
            content=ft.Column(
                [
                    ft.Text(f"Έκδοση: v{APP_VERSION}", size=15, color=ft.Colors.WHITE),
                    ft.Divider(height=10, color=ft.Colors.WHITE24),
                    ft.Text(
                        "Το Spyken μετατρέπει αρχεία .docx και .pdf σε αρχεία ήχου .mp3 "
                        "εναλλάσσοντας ανδρική και γυναικεία φωνή για κάθε παράγραφο, "
                        "καθώς και σε βίντεο .mp4 με οπτική ανάδειξη (highlight) "
                        "κάθε λέξης καθώς εκφωνείται.",
                        size=14, color=ft.Colors.GREY_300
                    ),
                    ft.Divider(height=10, color=ft.Colors.WHITE24),
                    ft.Text("🎶 Ελληνικές φωνές:", size=13, color=ft.Colors.BLUE_200),
                    ft.Text("  • Ανδρική: el-GR-NestorasNeural", size=12, color=ft.Colors.GREY_400),
                    ft.Text("  • Γυναικεία: el-GR-AthinaNeural", size=12, color=ft.Colors.GREY_400),
                    ft.Divider(height=5, color="transparent"),
                    ft.Text("🇬🇧 Αγγλικές φωνές (αυτόματη ανίχνευση):", size=13, color=ft.Colors.BLUE_200),
                    ft.Text("  • Ανδρική: en-GB-RyanNeural", size=12, color=ft.Colors.GREY_400),
                    ft.Text("  • Γυναικεία: en-GB-SoniaNeural", size=12, color=ft.Colors.GREY_400),
                    ft.Divider(height=10, color=ft.Colors.WHITE24),
                    ft.Text("by spyalekos • github.com/spyalekos", size=12, color=ft.Colors.GREY_500),
                ],
                tight=True,
            ),
            actions=[
                ft.TextButton("Κλείσιμο", on_click=lambda e: (setattr(dlg, 'open', False), page.update()))
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )
        page.overlay.append(dlg)
        dlg.open = True
        page.update()

    # ── Buttons ───────────────────────────────────────────────────────────────

    pick_btn = ft.ElevatedButton(
        "Επιλογή Αρχείων (.docx, .pdf)",
        icon=ft.Icons.FOLDER_OPEN,
        on_click=pick_files_clicked,
        style=ft.ButtonStyle(bgcolor=ft.Colors.INDIGO_700, color=ft.Colors.WHITE)
    )

    clear_btn = ft.ElevatedButton(
        "Καθαρισμός Λίστας",
        icon=ft.Icons.CLEAR,
        on_click=clear_files,
        style=ft.ButtonStyle(bgcolor=ft.Colors.RED_900, color=ft.Colors.WHITE)
    )

    convert_btn = ft.ElevatedButton(
        "Μετατροπή σε MP3",
        icon=ft.Icons.AUDIOTRACK,
        on_click=start_conversion,
        style=ft.ButtonStyle(bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE)
    )

    video_btn = ft.ElevatedButton(
        "Μετατροπή σε MP4",
        icon=ft.Icons.VIDEO_FILE,
        on_click=start_video_conversion,
        style=ft.ButtonStyle(bgcolor=ft.Colors.PURPLE_700, color=ft.Colors.WHITE)
    )

    # ── Layout ────────────────────────────────────────────────────────────────

    header = ft.Row(
        [
            ft.Icon(ft.Icons.AUDIOTRACK, size=40, color=ft.Colors.AMBER_400),
            ft.Text("Spyken", size=32, weight=ft.FontWeight.BOLD, color=ft.Colors.AMBER_400),
            ft.Text("by spyalekos  •  Έγγραφο σε MP3 & MP4", size=16, color=ft.Colors.GREY_300),
            ft.IconButton(
                icon=ft.Icons.INFO_OUTLINE,
                icon_color=ft.Colors.GREY_500,
                tooltip="Σχετικά",
                on_click=open_about,
            ),
        ],
        alignment=ft.MainAxisAlignment.CENTER,
    )

    button_row_top = ft.Row([pick_btn, clear_btn], alignment=ft.MainAxisAlignment.CENTER)
    button_row_bottom = ft.Row([convert_btn, video_btn], alignment=ft.MainAxisAlignment.CENTER, spacing=20)

    main_container = ft.Container(
        content=ft.Column(
            [
                header,
                ft.Divider(height=16, color="transparent"),
                button_row_top,
                ft.Container(
                    content=selected_files_list,
                    border=ft.border.all(1, ft.Colors.WHITE24),
                    border_radius=10,
                    padding=10,
                    expand=True
                ),
                ft.Divider(height=8, color="transparent"),
                button_row_bottom,
                ft.Divider(height=8, color="transparent"),
                ft.Column([status_text, progress_bar], horizontal_alignment=ft.CrossAxisAlignment.CENTER)
            ],
            horizontal_alignment=ft.CrossAxisAlignment.CENTER,
        ),
        expand=True
    )

    page.add(main_container)


if __name__ == "__main__":
    ft.app(target=main)
